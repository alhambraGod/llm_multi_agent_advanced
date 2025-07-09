"""
文档处理器 - 负责文档的解析、分块和预处理
"""

import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import hashlib
import asyncio
from concurrent.futures import ThreadPoolExecutor

# 文档处理库
import PyPDF2
import docx
from bs4 import BeautifulSoup
import markdown
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

# 多媒体处理
import cv2
import librosa
from PIL import Image
import pytesseract

class DocumentProcessor:
    """
    企业级文档处理器
    支持多种文档格式的解析、分块和预处理
    """
    
    def __init__(self, config: Dict[str, Any]):
        """
        初始化文档处理器
        
        Args:
            config: 配置参数
        """
        self.config = config
        self.logger = logging.getLogger(__name__)
        
        # 文本分割器配置
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.get('chunk_size', 1000),
            chunk_overlap=config.get('chunk_overlap', 200),
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
        
        # 支持的文件格式
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.txt': self._process_txt,
            '.md': self._process_markdown,
            '.html': self._process_html,
            '.jpg': self._process_image,
            '.jpeg': self._process_image,
            '.png': self._process_image,
            '.mp3': self._process_audio,
            '.wav': self._process_audio,
            '.mp4': self._process_video
        }
        
        # 线程池
        self.executor = ThreadPoolExecutor(max_workers=config.get('max_workers', 4))
    
    async def process_documents(self, file_paths: List[str]) -> List[Document]:
        """
        批量处理文档
        
        Args:
            file_paths: 文件路径列表
            
        Returns:
            处理后的文档列表
        """
        documents = []
        
        # 并发处理文档
        tasks = []
        for file_path in file_paths:
            task = asyncio.create_task(self._process_single_document(file_path))
            tasks.append(task)
        
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        for result in results:
            if isinstance(result, Exception):
                self.logger.error(f"文档处理失败: {result}")
            elif result:
                documents.extend(result)
        
        return documents
    
    async def _process_single_document(self, file_path: str) -> List[Document]:
        """
        处理单个文档
        
        Args:
            file_path: 文件路径
            
        Returns:
            文档块列表
        """
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                raise FileNotFoundError(f"文件不存在: {file_path}")
            
            # 获取文件扩展名
            ext = file_path.suffix.lower()
            if ext not in self.supported_formats:
                raise ValueError(f"不支持的文件格式: {ext}")
            
            # 处理文档
            processor = self.supported_formats[ext]
            content = await asyncio.get_event_loop().run_in_executor(
                self.executor, processor, str(file_path)
            )
            
            if not content:
                return []
            
            # 分块处理
            chunks = self.text_splitter.split_text(content)
            
            # 创建文档对象
            documents = []
            for i, chunk in enumerate(chunks):
                doc = Document(
                    page_content=chunk,
                    metadata={
                        'source': str(file_path),
                        'chunk_id': i,
                        'total_chunks': len(chunks),
                        'file_type': ext,
                        'file_hash': self._get_file_hash(str(file_path)),
                        'processed_at': asyncio.get_event_loop().time()
                    }
                )
                documents.append(doc)
            
            self.logger.info(f"成功处理文档: {file_path}, 生成 {len(documents)} 个块")
            return documents
            
        except Exception as e:
            self.logger.error(f"处理文档失败 {file_path}: {e}")
            raise
    
    def _process_pdf(self, file_path: str) -> str:
        """处理PDF文件"""
        content = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
        except Exception as e:
            self.logger.error(f"PDF处理失败: {e}")
            raise
        return content.strip()
    
    def _process_docx(self, file_path: str) -> str:
        """处理DOCX文件"""
        try:
            doc = docx.Document(file_path)
            content = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            return content.strip()
        except Exception as e:
            self.logger.error(f"DOCX处理失败: {e}")
            raise
    
    def _process_txt(self, file_path: str) -> str:
        """处理TXT文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            self.logger.error(f"TXT处理失败: {e}")
            raise
    
    def _process_markdown(self, file_path: str) -> str:
        """处理Markdown文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
                # 转换为HTML再提取文本
                html = markdown.markdown(md_content)
                soup = BeautifulSoup(html, 'html.parser')
                return soup.get_text().strip()
        except Exception as e:
            self.logger.error(f"Markdown处理失败: {e}")
            raise
    
    def _process_html(self, file_path: str) -> str:
        """处理HTML文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                html_content = file.read()
                soup = BeautifulSoup(html_content, 'html.parser')
                return soup.get_text().strip()
        except Exception as e:
            self.logger.error(f"HTML处理失败: {e}")
            raise
    
    def _process_image(self, file_path: str) -> str:
        """处理图像文件 - OCR文字识别"""
        try:
            image = Image.open(file_path)
            # 使用OCR提取文字
            text = pytesseract.image_to_string(image, lang='chi_sim+eng')
            return text.strip()
        except Exception as e:
            self.logger.error(f"图像处理失败: {e}")
            raise
    
    def _process_audio(self, file_path: str) -> str:
        """处理音频文件 - 语音转文字"""
        try:
            # 加载音频文件
            y, sr = librosa.load(file_path)
            
            # 这里应该集成语音识别服务（如Azure Speech, Google Speech等）
            # 暂时返回音频特征描述
            duration = librosa.get_duration(y=y, sr=sr)
            tempo, _ = librosa.beat.beat_track(y=y, sr=sr)
            
            return f"音频文件，时长: {duration:.2f}秒，节拍: {tempo:.2f} BPM"
        except Exception as e:
            self.logger.error(f"音频处理失败: {e}")
            raise
    
    def _process_video(self, file_path: str) -> str:
        """处理视频文件 - 提取关键帧和音频"""
        try:
            cap = cv2.VideoCapture(file_path)
            frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            fps = cap.get(cv2.CAP_PROP_FPS)
            duration = frame_count / fps
            
            cap.release()
            
            # 这里应该集成视频分析服务
            return f"视频文件，时长: {duration:.2f}秒，帧数: {frame_count}，帧率: {fps:.2f} FPS"
        except Exception as e:
            self.logger.error(f"视频处理失败: {e}")
            raise
    
    def _get_file_hash(self, file_path: str) -> str:
        """计算文件哈希值"""
        hash_md5 = hashlib.md5()
        try:
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except Exception as e:
            self.logger.error(f"计算文件哈希失败: {e}")
            return ""
    
    def __del__(self):
        """清理资源"""
        if hasattr(self, 'executor'):
            self.executor.shutdown(wait=True)